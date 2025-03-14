"""
Module xử lý trích xuất thông tin từ file Word
"""

import os
import re
import logging
import tempfile
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path

import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.core.logging import get_logger

logger = get_logger("cv_parser")


class WordParser:
    """
    Lớp trích xuất dữ liệu từ file Word
    """

    def __init__(self):
        """
        Khởi tạo parser
        """
        pass

    def extract_text(self, file_path: str) -> str:
        """
        Trích xuất văn bản từ file Word

        Args:
            file_path: Đường dẫn đến file Word

        Returns:
            str: Văn bản đã trích xuất
        """
        try:
            doc = docx.Document(file_path)
            full_text = []

            # Lặp qua tất cả các phần tử
            for element in self.iter_block_items(doc):
                if isinstance(element, Paragraph):
                    full_text.append(element.text)
                elif isinstance(element, Table):
                    # Xử lý bảng
                    table_text = self._process_table(element)
                    full_text.append(table_text)

            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Lỗi khi trích xuất văn bản từ Word {file_path}: {str(e)}")
            return ""

    def _process_table(self, table: Table) -> str:
        """
        Xử lý bảng trong văn bản Word

        Args:
            table: Đối tượng bảng từ python-docx

        Returns:
            str: Văn bản đã trích xuất từ bảng
        """
        table_text = []

        for i, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                # Lấy văn bản từ ô
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            # Thêm dòng vào bảng
            table_text.append(" | ".join(row_text))

        return "\n".join(table_text)

    def iter_block_items(self, parent):
        """
        Lặp qua tất cả các phần tử block (đoạn văn, bảng, v.v.)

        Args:
            parent: Đối tượng cha (thường là document)

        Yields:
            Paragraph hoặc Table
        """
        if isinstance(parent, Document):
            parent_element = parent.element.body
        elif isinstance(parent, _Cell):
            parent_element = parent._tc
        else:
            raise ValueError("Không thể lặp phần tử của đối tượng này")

        for child in parent_element.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Trích xuất metadata từ file Word

        Args:
            file_path: Đường dẫn đến file Word

        Returns:
            Dict[str, Any]: Metadata từ file Word
        """
        metadata = {}
        try:
            doc = docx.Document(file_path)
            core_properties = doc.core_properties

            # Lấy các thuộc tính cơ bản
            metadata_fields = [
                'author', 'category', 'comments', 'content_status',
                'created', 'identifier', 'keywords', 'language',
                'last_modified_by', 'last_printed', 'modified',
                'revision', 'subject', 'title', 'version'
            ]

            for field in metadata_fields:
                if hasattr(core_properties, field):
                    metadata[field] = getattr(core_properties, field)
        except Exception as e:
            logger.error(f"Lỗi khi trích xuất metadata: {str(e)}")

        return metadata

    def extract_images(self, file_path: str, output_dir: Optional[str] = None) -> List[str]:
        """
        Trích xuất hình ảnh từ file Word

        Args:
            file_path: Đường dẫn đến file Word
            output_dir: Thư mục đầu ra cho các hình ảnh

        Returns:
            List[str]: Danh sách đường dẫn đến các hình ảnh đã trích xuất
        """
        image_paths = []

        if output_dir is None:
            output_dir = tempfile.mkdtemp()
        else:
            os.makedirs(output_dir, exist_ok=True)

        try:
            doc = docx.Document(file_path)

            # Lấy các hình ảnh từ rels
            image_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image_extension = rel.target_ref.split(".")[-1]
                        image_path = os.path.join(output_dir, f'image_{image_index}.{image_extension}')

                        with open(image_path, 'wb') as f:
                            f.write(image_data)

                        image_paths.append(image_path)
                        image_index += 1
                    except Exception as e:
                        logger.error(f"Lỗi khi lưu hình ảnh: {str(e)}")

        except Exception as e:
            logger.error(f"Lỗi khi trích xuất hình ảnh: {str(e)}")

        return image_paths

    def get_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Lấy danh sách các đề mục trong văn bản

        Args:
            file_path: Đường dẫn đến file Word

        Returns:
            List[Dict[str, Any]]: Danh sách các đề mục
        """
        headings = []
        try:
            doc = docx.Document(file_path)

            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.replace('Heading ', ''))
                    headings.append({
                        'text': paragraph.text,
                        'level': level
                    })
        except Exception as e:
            logger.error(f"Lỗi khi lấy đề mục: {str(e)}")

        return headings

    def get_formatted_text(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Lấy văn bản có định dạng (in đậm, nghiêng, v.v.)

        Args:
            file_path: Đường dẫn đến file Word

        Returns:
            List[Dict[str, Any]]: Danh sách các đoạn văn bản và định dạng tương ứng
        """
        formatted_text = []
        try:
            doc = docx.Document(file_path)

            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue

                para_info = {
                    'text': paragraph.text,
                    'style': paragraph.style.name,
                    'runs': []
                }

                for run in paragraph.runs:
                    run_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font': run.font.name,
                        'size': run.font.size
                    }
                    para_info['runs'].append(run_info)

                formatted_text.append(para_info)

        except Exception as e:
            logger.error(f"Lỗi khi lấy văn bản định dạng: {str(e)}")

        return formatted_text